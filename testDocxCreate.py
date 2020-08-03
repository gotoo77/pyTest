from docx import Document
from docx.shared import Inches

document = Document()

document.add_picture("C:/Users/GAUTIERD1/PycharmProjects/test/assets/vr-gaming.png", width=Inches(1.25))

document.add_heading('Example facture', 0)

p = document.add_paragraph('Voici un exemple de paragraphe qui contient du ')
p.add_run('texte en gras').bold = True
p.add_run(' et aussi du ')
p.add_run('texte en italique.').italic = True

document.add_heading('Resume de titre (heading)', level=1)
document.add_paragraph('Citation démo', style='Intense Quote')

document.add_paragraph('liste 1')
document.add_paragraph('item 1 de la liste non ordonnée', style='List Bullet')
document.add_paragraph('item 2 de la liste non ordonnée', style='List Bullet')
document.add_paragraph('item 3 de la liste non ordonnée', style='List Bullet')
document.add_paragraph('liste 2')
document.add_paragraph('item 1 de la liste ordonnée', style='List Number')
document.add_paragraph('item 2 de la liste ordonnée', style='List Number')
document.add_paragraph('item 3 de la liste ordonnée', style='List Number')


records = (
    (3, '150', 'Redaction courrier'),
    (7, '300', 'conseil'),
    (4, '231', 'divers & recherches')
)

table = document.add_table(rows=1, cols=3)
hdr_cells = table.rows[0].cells
hdr_cells[0].text = 'Nb Heures'
hdr_cells[1].text = 'Montant (€)'
hdr_cells[2].text = 'Desc'
for qty, item_id, desc in records:
    row_cells = table.add_row().cells
    row_cells[0].text = str(qty)
    row_cells[1].text = item_id
    row_cells[2].text = desc

document.add_page_break()

document.save('demoFacture.docx')
